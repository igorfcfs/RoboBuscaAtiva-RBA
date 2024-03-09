import docx
import pandas as pd
import copy
import subprocess
import os
import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QPushButton, QVBoxLayout, QWidget

# Dicionário de padrões a serem substituídos
padroes = {
    '{{DIRETORIA}}': '',
    '{{ESCOLA}}': '',
    '{{RUA}}': '',
    '{{NUMERO}}': '',
    '{{BAIRRO}}': '',
    '{{CIDADE}}': '',
    '{{ESTADO}}': '',
    '{{CEP}}': '',
    '{{TELEFONE}}': '',
    '{{EMAIL}}': '',
}

# Funções originais para processar o documento Word
def processar_documento_word():
    # Carregue a planilha 'dados_filtrados_agregados'
    dados = pd.read_excel('ArquivosGerados/dados_filtrados_com_RA_e_Nome_e_Série.xlsx')

    # Abra o arquivo Word de modelo (documento_base)
    modelo = docx.Document('BasesDeDados/CONVOCACAO PARA COMPENSAR FALTAS.docx')

    # Crie um novo documento Word
    novo_doc = docx.Document()

    # Variável para rastrear o título anterior
    titulo_anterior = None

    # Itere sobre os dados da planilha
    for i, linha in dados.iterrows():
        nome = linha['Nome']
        numero = linha['RA']
        serie = linha['Série']

        # Copie o conteúdo do modelo para o novo documento com o mesmo estilo
        for elemento in modelo.element.body:
            novo_elemento = copy.deepcopy(elemento)
            novo_doc.element.body.append(novo_elemento)

        # Substitua apenas o padrão, mantendo o estilo
        for paragrafo in novo_doc.paragraphs:
            for run in paragrafo.runs:
                texto = run.text
                texto = texto.replace('XXXX', nome)
                texto = texto.replace('YYYY', str(numero))
                texto = texto.replace('ZZZZ', serie)
                run.text = texto

        # Adicione uma quebra de página após cada aluno, exceto o último
        if i < len(dados) - 1:
            novo_doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    # Salve o novo documento Word com as páginas duplicadas e as substituições
    novo_arquivo = 'ArquivosGerados/documento_convocacao.docx'
    novo_doc.save(novo_arquivo)

    print("Arquivo Word gerado com páginas separadas para cada aluno e as substituições de dados mantendo o estilo original.")

# Funções para adicionar quebras de página e recuperar conteúdo do arquivo gerado
def adicionar_quebras_de_pagina(doc):
    primeira_pagina = True  # Flag para controlar a primeira página
    for paragrafo in doc.paragraphs:
        if "GOVERNO DO ESTADO DE SÃO PAULO" in paragrafo.text:
            if primeira_pagina:
                primeira_pagina = False  # Desativa a flag na primeira ocorrência
            else:
                paragrafo.clear()
                run = paragrafo.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)
                run.add_text("GOVERNO DO ESTADO DE SÃO PAULO")

def remover_primeiro_paragrafo(doc):
    # Verifique se o documento tem parágrafos
    if len(doc.paragraphs) > 0:
        # Se houver parágrafos, exclua o primeiro
        primeiro_paragrafo = doc.paragraphs[0]
        doc.element.body.remove(primeiro_paragrafo._element)

def melhorar_conteudo_e_salvar():
    # Abra o arquivo
    arquivo_convocacao = 'ArquivosGerados/documento_convocacao.docx'

    # Adicionar quebra de páginas
    adicionar_quebras_de_pagina(arquivo_convocacao)

    # Remova o primeiro parágrafo (na primeira página)
    remover_primeiro_paragrafo(arquivo_convocacao)

    arquivo_convocacao.save(arquivo_convocacao)

    print(f'Conteúdo do arquivo recuperado e salvo em {arquivo_convocacao}')

# Função para executar o segundo script
def executar_script_converter_convocacao_para_pdf():
    # Obter o diretório do arquivo em execução
    diretorio_atual = os.path.dirname(os.path.abspath(__file__))

    # Nome do arquivo a ser executado (neste caso, na mesma pasta)
    caminho_segundo_script = "converter_convocacao_para_pdf.py"

    # Caminho completo para o segundo script
    caminho_completo = os.path.join(diretorio_atual, caminho_segundo_script)

    # Executar o segundo script
    subprocess.call(["python", caminho_completo])

# Classe principal da interface gráfica
class Formulario(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Formulário de Dados")
        self.setGeometry(100, 100, 400, 400)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        # Crie campos de entrada de texto para os dados
        for padrao in padroes:
            label = QLabel(f"Digite o valor para {padrao}:")
            input_field = QLineEdit()
            self.layout.addWidget(label)
            self.layout.addWidget(input_field)

            # Conecte o evento de edição ao método para atualizar os dados
            input_field.textChanged.connect(self.atualizar_dados)

        # Botão para realizar a substituição nos documentos
        self.botao_substituir = QPushButton("Substituir Padrões no Documento")
        self.botao_substituir.clicked.connect(self.substituir_no_documento)
        self.layout.addWidget(self.botao_substituir)

        self.central_widget.setLayout(self.layout)

    # Atualize o dicionário de padrões com os dados inseridos
    def atualizar_dados(self):
        for i, padrao in enumerate(padroes):
            padroes[padrao] = self.layout.itemAt(i * 2 + 1).widget().text()

    # Substitua os padrões no documento Word
    def substituir_no_documento(self):
        # Abra o arquivo Word de modelo
        modelo = docx.Document('BasesDeDados/CONVOCACAO PARA COMPENSAR FALTAS.docx')

        # Substitua os padrões no modelo
        for paragrafo in modelo.paragraphs:
            for padrao, substituicao in padroes.items():
                paragrafo.text = paragrafo.text.replace(padrao, substituicao)

        # Salve o novo documento Word
        novo_arquivo = 'ArquivosGerados/documento_com_dados_substituidos.docx'
        modelo.save(novo_arquivo)

        print(f'Documento gerado e salvo em {novo_arquivo}')

if __name__ == "__main__":
    # Primeiro, execute as funções para processar o documento Word
    processar_documento_word()

    # Em seguida, inicie a aplicação PyQt5
    app = QApplication(sys.argv)
    janela = Formulario()
    janela.show()
    sys.exit(app.exec_())
