import docx
import subprocess
import os
import sys
import pandas as pd
from PyQt5.QtWidgets import QApplication, QMainWindow, QLabel, QLineEdit, QPushButton, QVBoxLayout, QWidget

# Dicionário de padrões a serem substituídos
padroes = {
    '{{DIRETORIA}}': '',
    '{{ESCOLA}}': '',
    '{{RUA}}': '',
    '{{NUMERO}}': '',
    '{{BAIRRO}}': '',
    '{{CIDADE}}': '',
    '{{ESTADO}}': '',
    '{{CEP}}': '',
    '{{TELEFONE}}': '',
    '{{EMAIL}}': '',
}

# Função para salvar os dados preenchidos no dicionário de padrões
def salvar_dados():
    for padrao in padroes:
        padroes[padrao] = campo_de_texto[padrao].text()

# Função para processar o documento Word com base nos dados preenchidos
def processar_documento_word():
    # Carregue a planilha 'dados_filtrados_agregados'
    dados = pd.read_excel('ArquivosGerados/dados_filtrados_com_RA_e_Nome_e_Série.xlsx')

    # Abra o arquivo Word de modelo (documento_base)
    modelo = docx.Document('BasesDeDados/CONVOCACAO PARA COMPENSAR FALTAS.docx')

    # Crie um novo documento Word
    novo_doc = docx.Document()

    # Variável para rastrear o título anterior
    titulo_anterior = None

    # Itere sobre os dados da planilha
    for i, linha in dados.iterrows():
        nome = linha['Nome']
        numero = linha['RA']
        serie = linha['Série']

        # Copie o conteúdo do modelo para o novo documento com o mesmo estilo
        for elemento in modelo.element.body:
            novo_elemento = copy.deepcopy(elemento)
            novo_doc.element.body.append(novo_elemento)

        # Substitua apenas o padrão, mantendo o estilo
        for paragrafo in novo_doc.paragraphs:
            for run in paragrafo.runs:
                texto = run.text
                texto = texto.replace('XXXX', nome)
                texto = texto.replace('YYYY', str(numero))
                texto = texto.replace('ZZZZ', serie)
                run.text = texto

        # Adicione uma quebra de página após cada aluno, exceto o último
        if i < len(dados) - 1:
            novo_doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

    # Salve o novo documento Word com as páginas duplicadas e as substituições
    novo_arquivo = 'ArquivosGerados/documento_convocacao2.docx'
    novo_doc.save(novo_arquivo)

    print("Arquivo Word gerado com páginas separadas para cada aluno e as substituições de dados mantendo o estilo original.")

# Classe principal da interface gráfica
class Formulario(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Formulário de Dados")
        self.setGeometry(100, 100, 400, 400)

        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)

        self.layout = QVBoxLayout()

        # Crie campos de entrada de texto para os dados
        campo_de_texto = {}

        for padrao in padroes:
            label = QLabel(f"Digite o valor para {padrao}:")
            input_field = QLineEdit()
            campo_de_texto[padrao] = input_field
            self.layout.addWidget(label)
            self.layout.addWidget(input_field)

        # Botão para realizar a substituição nos documentos
        botao_substituir = QPushButton("Substituir Padrões no Documento")
        botao_substituir.clicked.connect(self.substituir_no_documento)
        self.layout.addWidget(botao_substituir)

        self.central_widget.setLayout(self.layout)

    # Substitua os padrões no documento Word
    def substituir_no_documento(self):
        # Salvar os dados preenchidos no dicionário de padrões
        salvar_dados()

        # Abra o arquivo Word de modelo
        modelo = docx.Document('BasesDeDados/CONVOCACAO PARA COMPENSAR FALTAS.docx')

        # Substitua os padrões no modelo
        for paragrafo in modelo.paragraphs:
            for padrao, substituicao in padroes.items():
                paragrafo.text = paragrafo.text.replace(padrao, substituicao)

        # Salve o novo documento Word
        novo_arquivo = 'ArquivosGerados/documento_com_dados_substituidos.docx'
        modelo.save(novo_arquivo)

        print(f'Documento gerado e salvo em {novo_arquivo}')

if __name__ == "__main__":
    # Inicie a aplicação PyQt5
    app = QApplication(sys.argv)
    janela = Formulario()
    janela.show()

    # Inicie o processamento do documento Word quando a aplicação for encerrada
    app.aboutToQuit.connect(processar_documento_word)

    sys.exit(app.exec_())
