import docx
import pandas as pd
import copy
import subprocess
import os

# Carregue a planilha 'dados_filtrados_agregados'
dados = pd.read_excel('ArquivosGerados/dados_filtrados_com_RA_e_Nome_e_Série.xlsx')

# Abra o arquivo Word de modelo (documento_base)
modelo = docx.Document('BasesDeDados/CONVOCACAO PARA COMPENSAR FALTAS.docx')  # Substitua 'documento_base.docx' pelo caminho do seu arquivo Word de modelo

# Crie um novo documento Word
novo_doc = docx.Document()

# Variável para rastrear o título anterior
titulo_anterior = None

# Itere sobre os dados da planilha
for i, linha in dados.iterrows():
    nome = linha['Nome']
    numero = linha['RA']
    serie = linha['Série']

    # Copie o conteúdo do modelo para o novo documento com o mesmo estilo
    for elemento in modelo.element.body:
        novo_elemento = copy.deepcopy(elemento)
        novo_doc.element.body.append(novo_elemento)

    # Substitua apenas o padrão, mantendo o estilo
    for paragrafo in novo_doc.paragraphs:
        for run in paragrafo.runs:
            texto = run.text
            texto = texto.replace('XXXX', nome)
            texto = texto.replace('YYYY', str(numero))
            texto = texto.replace('ZZZZ', serie)
            run.text = texto

    # Adicione uma quebra de página após cada aluno, exceto o último
    if i < len(dados) - 1:
        novo_doc.paragraphs[-1].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

# Salve o novo documento Word com as páginas duplicadas e as substituições
novo_arquivo = 'ArquivosGerados/documento_convocacao.docx'
novo_doc.save(novo_arquivo)

print("Arquivo Word gerado com páginas separadas para cada aluno e as substituições de dados mantendo o estilo original.")

# Função para adicionar quebras de página antes de cada ocorrência de "GOVERNO DO ESTADO DE SÃO PAULO"
def adicionar_quebras_de_pagina(doc):
    primeira_pagina = True  # Flag para controlar a primeira página
    for paragrafo in doc.paragraphs:
        if "GOVERNO DO ESTADO DE SÃO PAULO" in paragrafo.text:
            if primeira_pagina:
                primeira_pagina = False  # Desativa a flag na primeira ocorrência
            else:
                paragrafo.clear()
                run = paragrafo.add_run()
                run.add_break(docx.enum.text.WD_BREAK.PAGE)
                run.add_text("GOVERNO DO ESTADO DE SÃO PAULO")

def remover_primeiro_paragrafo(doc):
    # Verifique se o documento tem parágrafos
    if len(doc.paragraphs) > 0:
        # Se houver parágrafos, exclua o primeiro
        primeiro_paragrafo = doc.paragraphs[0]
        doc.element.body.remove(primeiro_paragrafo._element)

# Função para executar o segundo script
def executar_script_converter_convocacao_para_pdf():
    # Obter o diretório do arquivo em execução
    diretorio_atual = os.path.dirname(os.path.abspath(__file__))

    # Nome do arquivo a ser executado (neste caso, na mesma pasta)
    caminho_segundo_script = "converter_convocacao_para_pdf.py"

    # Caminho completo para o segundo script
    caminho_completo = os.path.join(diretorio_atual, caminho_segundo_script)

    # Executar o segundo script
    subprocess.call(["python", caminho_completo])

# Execute o segundo script
executar_script_converter_convocacao_para_pdf()

